"""Safely fill manuscript placeholders from completed formal-run manifests.

The default mode is a read-only dry run. Use ``--apply`` to write changes.
Every changed document is backed up before the first write, and unresolved
placeholders remain untouched.
"""

from __future__ import annotations

import argparse
import difflib
import json
import re
import shutil
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable


PLACEHOLDER_RE = re.compile(r"\{\{PENDING_[A-Z0-9_]+\}\}")
LATEX_PLACEHOLDER_RE = re.compile(
    r"\\\{\\\{PENDING(?:\\_[A-Z0-9]+)+\\\}\\\}"
)
RUN_PLACEHOLDER_RE = re.compile(
    r"\{\{PENDING_(R(?:\d{2}[AB]?)|S\d{2})_"
    r"(PRECISION|RECALL|MAP50|MAP75|MAP5095|PARAMS|GFLOPS|LATENCY|FPS)\}\}"
)
SUPPORTED_SUFFIXES = {".md", ".tex", ".docx"}


def _utc_stamp() -> str:
    return datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")


def _format_value(field: str, value: Any) -> str:
    if isinstance(value, bool) or not isinstance(value, (int, float)):
        raise TypeError(f"{field} must be numeric, got {type(value).__name__}")
    if field == "PARAMS":
        return str(int(value))
    if field in {"PRECISION", "RECALL", "MAP50", "MAP75", "MAP5095"}:
        return f"{float(value):.5f}"
    if field in {"GFLOPS", "LATENCY", "FPS"}:
        return f"{float(value):.4f}"
    raise KeyError(field)


def _manifest_values(path: Path) -> tuple[str, int, dict[str, str]]:
    payload = json.loads(path.read_text(encoding="utf-8"))
    run_id = str(payload.get("run_id", "")).upper()
    seed = int(payload.get("seed", -1))
    if not re.fullmatch(r"(?:R\d{2}[AB]?|S\d{2})", run_id):
        raise ValueError(f"Invalid run_id in {path}: {run_id!r}")
    if payload.get("status") != "completed":
        return run_id, seed, {}
    validation = payload.get("validation_metrics") or {}
    complexity = payload.get("complexity") or {}
    latency = complexity.get("latency") or {}
    raw = {
        "PRECISION": validation.get("precision"),
        "RECALL": validation.get("recall"),
        "MAP50": validation.get("map50"),
        "MAP75": validation.get("map75"),
        "MAP5095": validation.get("map50_95"),
        "PARAMS": complexity.get("parameters"),
        "GFLOPS": complexity.get("gflops"),
        "LATENCY": latency.get("mean_ms"),
        "FPS": latency.get("fps"),
    }
    values = {
        f"{{{{PENDING_{run_id}_{field}}}}}": _format_value(field, value)
        for field, value in raw.items()
        if value is not None
    }
    return run_id, seed, values


def collect_replacements(
    run_root: Path,
    *,
    seed: int = 0,
) -> tuple[dict[str, str], list[dict[str, Any]]]:
    """Collect unambiguous values from completed manifests at one seed."""

    replacements: dict[str, str] = {}
    provenance: list[dict[str, Any]] = []
    seen_runs: dict[str, Path] = {}
    for path in sorted(run_root.rglob("run_manifest.json")):
        run_id, manifest_seed, values = _manifest_values(path)
        if manifest_seed != seed or not values:
            continue
        previous = seen_runs.get(run_id)
        if previous is not None:
            _old_id, _old_seed, old_values = _manifest_values(previous)
            if old_values != values:
                raise ValueError(
                    f"Conflicting completed manifests for {run_id}: "
                    f"{previous} versus {path}"
                )
            continue
        seen_runs[run_id] = path
        replacements.update(values)
        provenance.append(
            {
                "run_id": run_id,
                "seed": manifest_seed,
                "manifest": str(path.resolve()),
                "tokens_available": sorted(values),
            }
        )
    return replacements, provenance


def _latex_placeholder(token: str) -> str:
    return (
        token.replace("{", r"\{")
        .replace("}", r"\}")
        .replace("_", r"\_")
    )


def _extract_text_placeholders(text: str) -> list[str]:
    raw = PLACEHOLDER_RE.findall(text)
    escaped = [
        value.replace(r"\{", "{")
        .replace(r"\}", "}")
        .replace(r"\_", "_")
        for value in LATEX_PLACEHOLDER_RE.findall(text)
    ]
    return raw + escaped


def _replace_text(text: str, replacements: dict[str, str]) -> tuple[str, Counter]:
    counts: Counter = Counter()
    tokens = sorted(
        set(_extract_text_placeholders(text)),
        key=len,
        reverse=True,
    )
    for token in tokens:
        value = replacements.get(token)
        if value is None:
            continue
        for serialized in (token, _latex_placeholder(token)):
            occurrences = text.count(serialized)
            if occurrences:
                text = text.replace(serialized, value)
                counts[token] += occurrences
    return text, counts


def _iter_table_paragraphs(table) -> Iterable:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs
            for nested in cell.tables:
                yield from _iter_table_paragraphs(nested)


def _iter_docx_paragraphs(document) -> Iterable:
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraphs(table)
    for section in document.sections:
        for container in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            yield from container.paragraphs
            for table in container.tables:
                yield from _iter_table_paragraphs(table)


def _replace_paragraph(paragraph, replacements: dict[str, str]) -> Counter:
    counts: Counter = Counter()
    while True:
        runs = list(paragraph.runs)
        full_text = "".join(run.text for run in runs)
        match = next(
            (
                match
                for match in PLACEHOLDER_RE.finditer(full_text)
                if match.group(0) in replacements
            ),
            None,
        )
        if match is None:
            break
        spans: list[tuple[int, int]] = []
        cursor = 0
        for run in runs:
            spans.append((cursor, cursor + len(run.text)))
            cursor += len(run.text)
        start_index = next(
            index
            for index, (_start, end) in enumerate(spans)
            if end > match.start()
        )
        end_index = next(
            index
            for index, (_start, end) in enumerate(spans)
            if end >= match.end()
        )
        start_offset = match.start() - spans[start_index][0]
        end_offset = match.end() - spans[end_index][0]
        token = match.group(0)
        value = replacements[token]
        if start_index == end_index:
            text = runs[start_index].text
            runs[start_index].text = (
                text[:start_offset] + value + text[end_offset:]
            )
        else:
            start_text = runs[start_index].text
            end_text = runs[end_index].text
            runs[start_index].text = start_text[:start_offset] + value
            for index in range(start_index + 1, end_index):
                runs[index].text = ""
            runs[end_index].text = end_text[end_offset:]
        counts[token] += 1
    return counts


def _docx_preview(path: Path, replacements: dict[str, str]) -> Counter:
    from docx import Document

    document = Document(path)
    counts: Counter = Counter()
    for paragraph in _iter_docx_paragraphs(document):
        text = "".join(run.text for run in paragraph.runs)
        for token in PLACEHOLDER_RE.findall(text):
            if token in replacements:
                counts[token] += 1
    return counts


def _apply_docx(path: Path, replacements: dict[str, str]) -> Counter:
    from docx import Document

    document = Document(path)
    counts: Counter = Counter()
    for paragraph in _iter_docx_paragraphs(document):
        counts.update(_replace_paragraph(paragraph, replacements))
    document.save(path)
    return counts


def _discover_documents(root: Path, explicit: list[str]) -> list[Path]:
    if explicit:
        documents = [Path(value).expanduser().resolve() for value in explicit]
    else:
        documents = sorted(
            path.resolve()
            for path in root.iterdir()
            if path.is_file() and path.suffix.lower() in SUPPORTED_SUFFIXES
        )
    unsupported = [
        path for path in documents if path.suffix.lower() not in SUPPORTED_SUFFIXES
    ]
    if unsupported:
        raise ValueError(f"Unsupported document(s): {unsupported}")
    missing = [path for path in documents if not path.is_file()]
    if missing:
        raise FileNotFoundError(missing)
    return documents


def update_documents(
    documents: list[Path],
    replacements: dict[str, str],
    *,
    apply: bool,
    backup_root: Path,
) -> tuple[list[dict[str, Any]], dict[str, str]]:
    records: list[dict[str, Any]] = []
    text_diffs: dict[str, str] = {}
    if apply:
        backup_root.mkdir(parents=True, exist_ok=False)
    for path in documents:
        suffix = path.suffix.lower()
        if suffix == ".docx":
            counts = _docx_preview(path, replacements)
            if apply and counts:
                shutil.copy2(path, backup_root / path.name)
                applied_counts = _apply_docx(path, replacements)
                if applied_counts != counts:
                    raise RuntimeError(f"DOCX replacement count changed for {path}")
        else:
            before = path.read_text(encoding="utf-8")
            after, counts = _replace_text(before, replacements)
            if counts:
                diff = "\n".join(
                    difflib.unified_diff(
                        before.splitlines(),
                        after.splitlines(),
                        fromfile=str(path),
                        tofile=str(path),
                        lineterm="",
                    )
                )
                text_diffs[str(path)] = diff
                if apply:
                    shutil.copy2(path, backup_root / path.name)
                    path.write_text(after, encoding="utf-8")
        records.append(
            {
                "path": str(path),
                "changed": bool(counts),
                "replacement_count": sum(counts.values()),
                "tokens": dict(sorted(counts.items())),
            }
        )
    return records, text_diffs


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "run_root",
        help="Root containing canonical run directories and run_manifest.json files.",
    )
    parser.add_argument(
        "--manuscript-root",
        default=r"D:\遥感船舶检测论文\01_论文初稿",
    )
    parser.add_argument(
        "--document",
        action="append",
        default=[],
        help="Explicit manuscript path; repeat as needed.",
    )
    parser.add_argument("--seed", type=int, default=0)
    mode = parser.add_mutually_exclusive_group()
    mode.add_argument("--apply", action="store_true")
    mode.add_argument("--dry-run", action="store_true")
    parser.add_argument("--report")
    args = parser.parse_args()

    run_root = Path(args.run_root).expanduser().resolve()
    manuscript_root = Path(args.manuscript_root).expanduser().resolve()
    if not run_root.is_dir():
        raise NotADirectoryError(run_root)
    if not manuscript_root.is_dir():
        raise NotADirectoryError(manuscript_root)
    documents = _discover_documents(manuscript_root, args.document)
    replacements, provenance = collect_replacements(run_root, seed=args.seed)
    stamp = _utc_stamp()
    backup_root = manuscript_root / "_result_backups" / stamp
    records, text_diffs = update_documents(
        documents,
        replacements,
        apply=args.apply,
        backup_root=backup_root,
    )
    report_path = (
        Path(args.report).expanduser().resolve()
        if args.report
        else manuscript_root / f"result_update_{stamp}.json"
    )
    unresolved: Counter = Counter()
    for path in documents:
        if path.suffix.lower() == ".docx":
            from docx import Document

            document = Document(path)
            for paragraph in _iter_docx_paragraphs(document):
                unresolved.update(PLACEHOLDER_RE.findall(paragraph.text))
        else:
            unresolved.update(
                _extract_text_placeholders(
                    path.read_text(encoding="utf-8")
                )
            )
    payload = {
        "schema_version": 1,
        "mode": "apply" if args.apply else "dry-run",
        "seed": args.seed,
        "run_root": str(run_root),
        "manuscript_root": str(manuscript_root),
        "completed_manifest_provenance": provenance,
        "available_replacements": dict(sorted(replacements.items())),
        "documents": records,
        "unresolved_placeholders_after_mode": dict(sorted(unresolved.items())),
        "backup_directory": str(backup_root) if args.apply else None,
        "text_diffs": text_diffs,
    }
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(
        json.dumps(payload, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(json.dumps(payload, ensure_ascii=False, indent=2))
    print(f"Report: {report_path}")


if __name__ == "__main__":
    main()

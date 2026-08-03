"""Build bilingual editable experiment-chapter skeletons with placeholders."""

from __future__ import annotations

import argparse
import json
import os
from pathlib import Path
from typing import Any

RUNS = [
    "R00",
    "R01",
    "R02",
    "R03",
    "R04",
    "R05A",
    "R05B",
    "R06",
    "R07",
    "R08",
    "R09",
    "R10",
    "R11",
    "R12",
    "R13",
    "S00",
    "S01",
]


def p(run: str, field: str) -> str:
    return "{{PENDING_" + run + "_" + field + "}}"


def pending(name: str) -> str:
    return "{{PENDING_" + name + "}}"


TABLES: dict[int, dict[str, Any]] = {
    1: {
        "zh": "数据集统计",
        "en": "Dataset statistics",
        "headers_zh": ["数据集", "划分", "图像数", "实例数", "空标签图像", "来源/许可"],
        "headers_en": ["Dataset", "Split", "Images", "Instances", "Empty-label images", "Source/license"],
        "rows": [
            ["Primary", "Train", pending("PRIMARY_TRAIN_IMAGES"), pending("PRIMARY_TRAIN_INSTANCES"), pending("PRIMARY_TRAIN_EMPTY"), pending("PRIMARY_SOURCE_LICENSE")],
            ["Primary", "Validation", pending("PRIMARY_VAL_IMAGES"), pending("PRIMARY_VAL_INSTANCES"), pending("PRIMARY_VAL_EMPTY"), pending("PRIMARY_SOURCE_LICENSE")],
            ["Primary", "Test", pending("PRIMARY_TEST_IMAGES"), pending("PRIMARY_TEST_INSTANCES"), pending("PRIMARY_TEST_EMPTY"), pending("PRIMARY_SOURCE_LICENSE")],
            [pending("SECOND_DATASET_NAME"), "Train/Validation/Test", pending("SECOND_DATASET_SPLITS"), pending("SECOND_DATASET_INSTANCES"), pending("SECOND_DATASET_EMPTY"), pending("SECOND_DATASET_LICENSE")],
        ],
    },
    2: {
        "zh": "正式实验设置",
        "en": "Formal experimental settings",
        "headers_zh": ["项目", "冻结值/记录方式"],
        "headers_en": ["Item", "Frozen value / recording policy"],
        "rows": [
            ["Python", "3.12.x"],
            ["PyTorch / CUDA / GPU", f"{pending('FORMAL_TORCH_VERSION')} / {pending('FORMAL_CUDA_VERSION')} / {pending('FORMAL_GPU')}"],
            ["Ultralytics", "8.4.92"],
            ["Input / epochs / batch / workers", "640 / 150 / 8 / 2"],
            ["Optimizer policy", "auto (effective choice retained in args.yaml and run manifest)"],
            ["lr0 / weight decay", "0.01 / 0.0005"],
            ["Seed", "0; stability runs use 0, 1, 2 for R00/R02/R10"],
            ["Initialization", "yolo11n.pt, yolo11s.pt, or yolov8n.pt; no preceding ablation checkpoint"],
            ["Model selection", "highest validation mAP50-95; test split excluded"],
        ],
    },
    3: {
        "zh": "与已有方法的比较",
        "en": "Comparison with existing methods",
        "headers_zh": ["方法", "年份", "输入", "P", "R", "mAP50", "mAP50-95", "Params", "GFLOPs", "FPS", "协议"],
        "headers_en": ["Method", "Year", "Input", "P", "R", "mAP50", "mAP50-95", "Params", "GFLOPs", "FPS", "Protocol"],
        "rows": [[pending("SOTA_METHOD_1"), pending("SOTA_YEAR_1"), pending("SOTA_INPUT_1"), pending("SOTA_P_1"), pending("SOTA_R_1"), pending("SOTA_MAP50_1"), pending("SOTA_MAP5095_1"), pending("SOTA_PARAMS_1"), pending("SOTA_GFLOPS_1"), pending("SOTA_FPS_1"), pending("SOTA_PROTOCOL_1")]],
    },
    4: {
        "zh": "累加消融",
        "en": "Cumulative ablation",
        "headers_zh": ["ID", "DPLS", "CA-SCAM", "VGUP", "P", "R", "mAP50", "mAP50-95", "Params", "GFLOPs"],
        "headers_en": ["ID", "DPLS", "CA-SCAM", "VGUP", "P", "R", "mAP50", "mAP50-95", "Params", "GFLOPs"],
        "rows": [
            ["A0", "No", "No", "No", p("R00", "PRECISION"), p("R00", "RECALL"), p("R00", "MAP50"), p("R00", "MAP5095"), p("R00", "PARAMS"), p("R00", "GFLOPS")],
            ["A1", "Yes", "No", "No", p("R02", "PRECISION"), p("R02", "RECALL"), p("R02", "MAP50"), p("R02", "MAP5095"), p("R02", "PARAMS"), p("R02", "GFLOPS")],
            ["A2", "Yes", "Yes", "No", p("R04", "PRECISION"), p("R04", "RECALL"), p("R04", "MAP50"), p("R04", "MAP5095"), p("R04", "PARAMS"), p("R04", "GFLOPS")],
            ["A3", "Yes", "Yes", "Yes", p("R10", "PRECISION"), p("R10", "RECALL"), p("R10", "MAP50"), p("R10", "MAP5095"), p("R10", "PARAMS"), p("R10", "GFLOPS")],
        ],
    },
    5: {
        "zh": "DPLS 专项消融",
        "en": "DPLS controlled ablation",
        "headers_zh": ["ID", "金字塔", "上采样", "P", "R", "mAP50", "mAP50-95", "Params", "GFLOPs"],
        "headers_en": ["ID", "Pyramid", "Upsampling", "P", "R", "mAP50", "mAP50-95", "Params", "GFLOPs"],
        "rows": [
            ["D0", "P3-P5", "Nearest", p("R00", "PRECISION"), p("R00", "RECALL"), p("R00", "MAP50"), p("R00", "MAP5095"), p("R00", "PARAMS"), p("R00", "GFLOPS")],
            ["D1", "P2-P4", "Nearest", p("R01", "PRECISION"), p("R01", "RECALL"), p("R01", "MAP50"), p("R01", "MAP5095"), p("R01", "PARAMS"), p("R01", "GFLOPS")],
            ["D2", "P2-P4", "DySample", p("R02", "PRECISION"), p("R02", "RECALL"), p("R02", "MAP50"), p("R02", "MAP5095"), p("R02", "PARAMS"), p("R02", "GFLOPS")],
        ],
    },
    6: {
        "zh": "SCAM 与 CA-SCAM 消融",
        "en": "SCAM and CA-SCAM ablations",
        "headers_zh": ["ID", "变体", "P", "R", "mAP50", "mAP50-95", "Params"],
        "headers_en": ["ID", "Variant", "P", "R", "mAP50", "mAP50-95", "Params"],
        "rows": [
            ["C0", "No attention", p("R02", "PRECISION"), p("R02", "RECALL"), p("R02", "MAP50"), p("R02", "MAP5095"), p("R02", "PARAMS")],
            ["C1/CI0", "Original SCAM", p("R03", "PRECISION"), p("R03", "RECALL"), p("R03", "MAP50"), p("R03", "MAP5095"), p("R03", "PARAMS")],
            ["CI1", "Contrast map + fixed beta", p("R05A", "PRECISION"), p("R05A", "RECALL"), p("R05A", "MAP50"), p("R05A", "MAP5095"), p("R05A", "PARAMS")],
            ["CI2", "Contrast map + learnable unbounded beta", p("R05B", "PRECISION"), p("R05B", "RECALL"), p("R05B", "MAP50"), p("R05B", "MAP5095"), p("R05B", "PARAMS")],
            ["C2/CI3", "Complete bounded CA-SCAM", p("R04", "PRECISION"), p("R04", "RECALL"), p("R04", "MAP50"), p("R04", "MAP5095"), p("R04", "PARAMS")],
        ],
    },
    7: {
        "zh": "ERUP 与 VGUP 比较",
        "en": "ERUP versus VGUP",
        "headers_zh": ["ID", "输入处理", "P", "R", "mAP50", "mAP50-95", "模块 Params", "总 Params", "GFLOPs", "延迟"],
        "headers_en": ["ID", "Input processor", "P", "R", "mAP50", "mAP50-95", "Module params", "Total params", "GFLOPs", "Latency"],
        "rows": [
            ["V0", "None", p("R04", "PRECISION"), p("R04", "RECALL"), p("R04", "MAP50"), p("R04", "MAP5095"), "0", p("R04", "PARAMS"), p("R04", "GFLOPS"), p("R04", "LATENCY")],
            ["V1", "ERUP", p("R06", "PRECISION"), p("R06", "RECALL"), p("R06", "MAP50"), p("R06", "MAP5095"), "6,781,042", p("R06", "PARAMS"), p("R06", "GFLOPS"), p("R06", "LATENCY")],
            ["V2", "VGUP", p("R10", "PRECISION"), p("R10", "RECALL"), p("R10", "MAP50"), p("R10", "MAP5095"), "77,396", p("R10", "PARAMS"), p("R10", "GFLOPS"), p("R10", "LATENCY")],
        ],
    },
    8: {
        "zh": "VGUP 门控消融",
        "en": "VGUP gate ablation",
        "headers_zh": ["ID", "全局门", "空间门", "P", "R", "mAP50", "mAP50-95", "Params"],
        "headers_en": ["ID", "Global gate", "Spatial gate", "P", "R", "mAP50", "mAP50-95", "Params"],
        "rows": [
            ["VG0", "No", "No", p("R07", "PRECISION"), p("R07", "RECALL"), p("R07", "MAP50"), p("R07", "MAP5095"), p("R07", "PARAMS")],
            ["VG1", "Yes", "No", p("R08", "PRECISION"), p("R08", "RECALL"), p("R08", "MAP50"), p("R08", "MAP5095"), p("R08", "PARAMS")],
            ["VG2", "No", "Yes", p("R09", "PRECISION"), p("R09", "RECALL"), p("R09", "MAP50"), p("R09", "MAP5095"), p("R09", "PARAMS")],
            ["VG3", "Yes", "Yes", p("R10", "PRECISION"), p("R10", "RECALL"), p("R10", "MAP50"), p("R10", "MAP5095"), p("R10", "PARAMS")],
        ],
    },
    9: {
        "zh": "第二数据集独立训练结果",
        "en": "Independent second-dataset results",
        "headers_zh": ["ID", "方法", "P", "R", "mAP50", "mAP50-95", "Params", "GFLOPs"],
        "headers_en": ["ID", "Method", "P", "R", "mAP50", "mAP50-95", "Params", "GFLOPs"],
        "rows": [
            ["S0", "YOLO11n baseline", p("S00", "PRECISION"), p("S00", "RECALL"), p("S00", "MAP50"), p("S00", "MAP5095"), p("S00", "PARAMS"), p("S00", "GFLOPS")],
            ["S1", "YOLO11n final", p("S01", "PRECISION"), p("S01", "RECALL"), p("S01", "MAP50"), p("S01", "MAP5095"), p("S01", "PARAMS"), p("S01", "GFLOPS")],
        ],
    },
    10: {
        "zh": "跨模型泛化",
        "en": "Cross-model generalization",
        "headers_zh": ["检测器", "Baseline/Final", "P", "R", "mAP50", "mAP50-95", "Params", "GFLOPs"],
        "headers_en": ["Detector", "Baseline/Final", "P", "R", "mAP50", "mAP50-95", "Params", "GFLOPs"],
        "rows": [
            ["YOLO11n", "Baseline", p("R00", "PRECISION"), p("R00", "RECALL"), p("R00", "MAP50"), p("R00", "MAP5095"), p("R00", "PARAMS"), p("R00", "GFLOPS")],
            ["YOLO11n", "Final", p("R10", "PRECISION"), p("R10", "RECALL"), p("R10", "MAP50"), p("R10", "MAP5095"), p("R10", "PARAMS"), p("R10", "GFLOPS")],
            ["YOLO11s", "Baseline", p("R13", "PRECISION"), p("R13", "RECALL"), p("R13", "MAP50"), p("R13", "MAP5095"), p("R13", "PARAMS"), p("R13", "GFLOPS")],
            ["YOLOv8n", "Baseline", p("R11", "PRECISION"), p("R11", "RECALL"), p("R11", "MAP50"), p("R11", "MAP5095"), p("R11", "PARAMS"), p("R11", "GFLOPS")],
            ["YOLOv8n", "Final", p("R12", "PRECISION"), p("R12", "RECALL"), p("R12", "MAP50"), p("R12", "MAP5095"), p("R12", "PARAMS"), p("R12", "GFLOPS")],
        ],
    },
    11: {
        "zh": "复杂度与推理效率",
        "en": "Complexity and inference efficiency",
        "headers_zh": ["方法", "输入", "Params", "GFLOPs", "FP32 延迟", "FP16 延迟", "FPS", "峰值显存", "mAP50-95"],
        "headers_en": ["Method", "Input", "Params", "GFLOPs", "FP32 latency", "FP16 latency", "FPS", "Peak memory", "mAP50-95"],
        "rows": [
            ["YOLO11n baseline", "640", p("R00", "PARAMS"), p("R00", "GFLOPS"), p("R00", "LATENCY"), pending("R00_FP16_LATENCY"), p("R00", "FPS"), pending("R00_PEAK_MEMORY"), p("R00", "MAP5095")],
            ["YOLO11n final", "640", p("R10", "PARAMS"), p("R10", "GFLOPS"), p("R10", "LATENCY"), pending("R10_FP16_LATENCY"), p("R10", "FPS"), pending("R10_PEAK_MEMORY"), p("R10", "MAP5095")],
            ["YOLO11s baseline", "640", p("R13", "PARAMS"), p("R13", "GFLOPS"), p("R13", "LATENCY"), pending("R13_FP16_LATENCY"), p("R13", "FPS"), pending("R13_PEAK_MEMORY"), p("R13", "MAP5095")],
        ],
    },
    12: {
        "zh": "多随机种子稳定性",
        "en": "Multi-seed stability",
        "headers_zh": ["模型", "Seeds", "Recall (mean ± std)", "mAP50 (mean ± std)", "mAP50-95 (mean ± std)"],
        "headers_en": ["Model", "Seeds", "Recall (mean ± std)", "mAP50 (mean ± std)", "mAP50-95 (mean ± std)"],
        "rows": [
            ["Baseline / R00", "0,1,2", pending("R00_RECALL_MEAN_STD"), pending("R00_MAP50_MEAN_STD"), pending("R00_MAP5095_MEAN_STD")],
            ["DPLS / R02", "0,1,2", pending("R02_RECALL_MEAN_STD"), pending("R02_MAP50_MEAN_STD"), pending("R02_MAP5095_MEAN_STD")],
            ["Final / R10", "0,1,2", pending("R10_RECALL_MEAN_STD"), pending("R10_MAP50_MEAN_STD"), pending("R10_MAP5095_MEAN_STD")],
        ],
    },
}


SECTIONS = [
    {
        "number": "4.1",
        "zh": "数据集与数据预处理",
        "en": "Dataset and Data Preprocessing",
        "zh_paragraphs": [
            f"主数据集为 {pending('PRIMARY_DATASET_NAME')}，来源与许可信息为 {pending('PRIMARY_DATASET_SOURCE')}。正式实验严格沿用冻结后的 train/validation/test 划分，不重新清洗、不重新分配样本，也不以测试集选择模型。图像和标签仅从 Google Drive 只读复制到 Colab 本地目录，以减少训练 I/O；复制过程同时记录文件数、字节数和 SHA256 审计信息。",
            f"标签采用 YOLO 检测格式，类别映射为 {pending('PRIMARY_CLASS_MAPPING')}。所有图像通过保持宽高比的 letterbox 变换统一到 640×640。训练增强遵循 Ultralytics 8.4.92 的冻结配置；验证集和测试集关闭增强，且不混入训练样本。空标签图像被保留为合法负样本，其数量与实例统计将在表 1 中由正式审计文件填入。",
            f"第二数据集尚未确定，暂以 {pending('SECOND_DATASET_NAME')} 占位。选定后须补齐许可、引用、类别映射、划分、异常框和跨划分重复审计，并从各自官方预训练权重独立训练 S0/S1，不能加载主数据集最终权重作为默认微调起点。",
        ],
        "en_paragraphs": [
            f"The primary dataset is {pending('PRIMARY_DATASET_NAME')}, with provenance and license recorded as {pending('PRIMARY_DATASET_SOURCE')}. All formal experiments preserve the frozen train/validation/test split. No sample is cleaned or repartitioned, and the test split is excluded from checkpoint selection. Images and labels are copied read-only from Google Drive to the Colab local disk; the transfer records file counts, processed bytes, and checksum evidence.",
            f"Annotations follow the YOLO detection format with class mapping {pending('PRIMARY_CLASS_MAPPING')}. Aspect-ratio-preserving letterbox resizing produces a common 640×640 input. Training augmentation follows the frozen Ultralytics 8.4.92 configuration, whereas validation and test data are not augmented and never receive training samples. Empty-label images remain valid negatives; their counts and instance statistics will be filled from the formal audit in Table 1.",
            f"The second dataset has not yet been selected and is represented by {pending('SECOND_DATASET_NAME')}. After selection, its license, citation, class mapping, split integrity, abnormal boxes, and cross-split duplication must be audited. S0 and S1 will be trained independently from the appropriate official pretrained checkpoint rather than by default fine-tuning the primary-dataset final model.",
        ],
        "tables": [1],
        "figures": [
            (1, "典型海洋遥感场景示例", "Representative maritime remote-sensing scenes"),
            (2, "船舶短边、长边、面积和宽高比分布", "Distributions of ship short side, long side, area, and aspect ratio"),
        ],
    },
    {
        "number": "4.2",
        "zh": "实验设置",
        "en": "Experimental Setup",
        "zh_paragraphs": [
            "正式模型在 Google Colab GPU 环境中训练，本地 Windows 仅执行模型构建、CPU 前向/反向、权重继承与测试。所有唯一结构采用同一输入尺寸、epoch、batch、优化器策略、增强和模型选择规则。R00/R02/R10 预留 seeds 0、1、2 的稳定性实验，其余结构先以 seed 0 完成。",
            "每个 YOLO11n 结构从官方 yolo11n.pt 独立初始化，YOLOv8n 结构从 yolov8n.pt 独立初始化。层级变化采用显式语义层映射和同形状检查，训练器在 epoch 1 前进行逐张量交接审计；严禁使用前一消融实验的 best.pt 继续训练。",
        ],
        "en_paragraphs": [
            "Formal models are trained on a Google Colab GPU, while the Windows workstation is restricted to model construction, CPU forward/backward checks, inheritance audits, and tests. Every unique topology uses the same input size, epoch budget, batch size, optimizer policy, augmentation, and model-selection rule. Stability runs with seeds 0, 1, and 2 are reserved for R00, R02, and R10; the other variants initially use seed 0.",
            "Each YOLO11n topology is independently initialized from official yolo11n.pt, and each YOLOv8n topology from yolov8n.pt. Structural index changes are handled by explicit semantic layer mapping and shape checks. Before epoch 1, a tensor-by-tensor trainer handoff audit verifies the initialized state. Continuing from the best checkpoint of a preceding ablation is prohibited.",
        ],
        "tables": [2],
        "figures": [],
    },
    {
        "number": "4.3",
        "zh": "评价指标",
        "en": "Evaluation Metrics",
        "zh_paragraphs": [
            "检测性能采用 Precision、Recall、mAP50 和 mAP50-95。模型效率采用参数量、GFLOPs、模型大小、PyTorch 前向延迟、FPS 和峰值 GPU 显存。速度比较仅在 GPU、输入、batch、精度、框架以及是否包含预处理/NMS 完全一致时进行。",
            "分尺度评价基于 640 letterbox 后的目标短边和面积，报告每组实例数、TP、FP、FN、Precision 与 Recall。若分组 AP 未通过独立实现验证，则不以分组 Precision/Recall 的简单平均冒充 AP 或整体 mAP。",
        ],
        "en_paragraphs": [
            "Detection performance is measured by Precision, Recall, mAP50, and mAP50-95. Efficiency is characterized by parameter count, GFLOPs, model size, PyTorch forward latency, FPS, and peak GPU memory. Speed values are directly compared only when GPU, input size, batch size, numerical precision, framework, preprocessing scope, and NMS scope are identical.",
            "Scale-stratified evaluation uses target short side and area after 640 letterbox resizing and reports the number of instances together with TP, FP, FN, Precision, and Recall. Unless grouped AP has been independently validated, a simple mean of grouped Precision/Recall is not presented as AP or overall mAP.",
        ],
        "formulas": [
            r"P = TP / (TP + FP)",
            r"R = TP / (TP + FN)",
            r"AP = \int_0^1 P(R)\,dR",
            r"mAP_{50:95} = (1/10)\sum_{\tau=0.50}^{0.95} AP_\tau",
        ],
        "tables": [],
        "figures": [],
    },
    {
        "number": "4.4",
        "zh": "与已有先进方法的比较",
        "en": "Comparison with State-of-the-Art Methods",
        "zh_paragraphs": [
            "本节仅预留统一重训与外部论文报告结果的接口。二者将在表 3 中明确标注训练协议、数据集、输入尺寸和速度测试口径；不同数据集或不同硬件下的数值不作无条件横向排序。",
            f"正式结果填入后，再依据证据讨论本文方法与比较模型的差异。目前不预写优越性结论：{pending('SOTA_INTERPRETATION')}",
        ],
        "en_paragraphs": [
            "This section reserves separate interfaces for uniformly retrained models and externally reported results. Table 3 will explicitly label the training protocol, dataset, input size, and speed-measurement scope. Values obtained on different datasets or hardware will not be ranked without qualification.",
            f"Evidence-based interpretation will be written only after the formal results are available. No superiority claim is pre-authored here: {pending('SOTA_INTERPRETATION')}",
        ],
        "tables": [3],
        "figures": [],
    },
    {
        "number": "4.5",
        "zh": "累加消融实验",
        "en": "Cumulative Ablation Study",
        "zh_paragraphs": [
            "累加消融由 A0（YOLO11n）、A1（+DPLS）、A2（+CA-SCAM）和 A3（+VGUP）组成，不包含 InceptionDW。A1、A2 和 A3 均从官方 yolo11n.pt 独立初始化，因而表 4 比较的是结构差异，而不是连续微调收益。",
            f"表 4 的定量解释在正式结果填入后完成。当前仅保留中性占位：{pending('CUMULATIVE_ABLATION_INTERPRETATION')}",
        ],
        "en_paragraphs": [
            "The cumulative study comprises A0 (YOLO11n), A1 (+DPLS), A2 (+CA-SCAM), and A3 (+VGUP); InceptionDW is not included. A1, A2, and A3 are independently initialized from official yolo11n.pt, so Table 4 compares topologies rather than gains from sequential fine-tuning.",
            f"Quantitative interpretation will be completed after formal values are inserted. The current neutral placeholder is {pending('CUMULATIVE_ABLATION_INTERPRETATION')}.",
        ],
        "tables": [4],
        "figures": [],
    },
    {
        "number": "4.6",
        "zh": "DPLS 分析",
        "en": "Analysis of DPLS",
        "zh_paragraphs": [
            "现有确定性尺度分析用于描述 640 letterbox 后船舶短边、长边、面积和宽高比分布，并计算 P2/P3/P4/P5 的空间稀释率。该统计只构成尺度动机，不预先证明 DPLS 必然有效；结构有效性由 D0/D1/D2 的统一训练结果验证。",
            "D0 为官方 P3/P4/P5 检测；D1 将金字塔前移至 P2/P3/P4 并保留 nearest 上采样；D2 仅将 D1 的两个上采样替换为 DySample。因此 D0→D1 分离 PLS 整体贡献，D1→D2 分离动态上采样贡献。",
            f"分短边、分面积与检测层预 NMS 贡献结果将分别填入 {pending('DPLS_SCALE_RESULTS')}、{pending('DPLS_LEVEL_RESULTS')}。稀疏组会由脚本给出合并建议并同时报告实例数。",
        ],
        "en_paragraphs": [
            "The deterministic scale analysis describes ship short side, long side, area, and aspect ratio after 640 letterbox resizing and calculates spatial dilution at P2/P3/P4/P5. These statistics motivate the scale design but do not predetermine that DPLS is effective; effectiveness is tested by uniformly trained D0/D1/D2 models.",
            "D0 is the official P3/P4/P5 detector. D1 shifts the pyramid to P2/P3/P4 while retaining nearest-neighbor upsampling. D2 changes only the two D1 upsamplers to DySample. Consequently, D0→D1 isolates the complete PLS shift, whereas D1→D2 isolates dynamic upsampling.",
            f"Short-side, area, and pre-NMS level-contribution results will be inserted at {pending('DPLS_SCALE_RESULTS')} and {pending('DPLS_LEVEL_RESULTS')}. Sparse bins receive an explicit merge recommendation and always report instance counts.",
        ],
        "tables": [5],
        "figures": [
            (3, "D0/D1/D2 同图预测与 P2-P5 真实特征响应", "Matched D0/D1/D2 predictions and real P2-P5 feature responses"),
        ],
    },
    {
        "number": "4.7",
        "zh": "CA-SCAM 分析",
        "en": "Analysis of CA-SCAM",
        "zh_paragraphs": [
            "SCAM 是已有上下文注意方法；本文的 CA-SCAM 在其上下文残差上增加局部对比度估计、3×3 空间投影以及残差幅度校准。完整形式采用 beta=0.1×tanh(logit)，从而将校准幅度限制在 ±0.1。",
            "CI0 为原始 SCAM；CI1 使用固定 beta；CI2 使用零初始化、可学习但无界的 beta；CI3 为可学习有界 beta。CI2 与 CI3 在初始化时均与原始 SCAM 精确等价。该序列分别验证空间对比度校准、强度学习和有界约束。",
            f"输入特征、原 SCAM 残差、局部对比度图、空间校准图、校准后残差和输出特征均由真实 forward hook 生成。beta 轨迹和定量结论占位为 {pending('CA_SCAM_BETA_ANALYSIS')}。",
        ],
        "en_paragraphs": [
            "SCAM is an existing contextual attention method. CA-SCAM augments its context residual with local-contrast estimation, a 3×3 spatial projection, and residual-amplitude calibration. The complete formulation uses beta=0.1×tanh(logit), restricting the calibration amplitude to ±0.1.",
            "CI0 is original SCAM; CI1 uses a fixed beta; CI2 uses a zero-initialized, learnable but unbounded beta; and CI3 uses a learnable bounded beta. CI2 and CI3 are exactly equivalent to SCAM at initialization. This sequence isolates spatial contrast calibration, strength learning, and the bounding constraint.",
            f"Input features, original SCAM residuals, local-contrast maps, spatial calibration maps, calibrated residuals, and output features are generated from real forward hooks. The beta trajectory and quantitative interpretation remain at {pending('CA_SCAM_BETA_ANALYSIS')}.",
        ],
        "tables": [6],
        "figures": [
            (4, "P2/P3/P4 的 CA-SCAM 真实前向校准过程", "Real CA-SCAM forward calibration at P2/P3/P4"),
        ],
    },
    {
        "number": "4.8",
        "zh": "VGUP 分析",
        "en": "Analysis of VGUP",
        "zh_paragraphs": [
            "VGUP 位于模型输入端，保留 ERUP 的 BPW 与 KBL 可微图像处理思想，并使用轻量编码器同时预测滤波参数和两类门控。全局标量门控制 BPW 残差接受程度，空间门逐像素控制 KBL 残差。",
            "当前代码重新计算得到 ERUP 为 6,781,042 个可训练参数，完整 VGUP 为 77,396 个，VGUP/ERUP=0.0114135851（1.1414%）。该比例来自当前正式实现，不使用早期“约 1/8”的估计。",
            "VG0-VG3 去门时仍保留 BPW/KBL 算子，并令对应有效门值为 1；因此只比较门控机制。全局门直方图、空间门均值/标准差、亮度与对比度关系及过度处理失败案例均由完整验证集统计与人工复核共同形成。",
            f"性能与门控行为的正式解释占位为 {pending('VGUP_INTERPRETATION')}。",
        ],
        "en_paragraphs": [
            "VGUP is placed at the model input. It retains the differentiable BPW and KBL processing concept of ERUP while using a lightweight encoder to predict filter parameters and two gates. A global scalar gate controls acceptance of the BPW residual, whereas a spatial gate controls the KBL residual pixel by pixel.",
            "Recomputation from the current code yields 6,781,042 trainable parameters for ERUP and 77,396 for complete VGUP, giving VGUP/ERUP=0.0114135851 (1.1414%). This ratio is derived from the formal implementation and replaces the earlier approximate one-eighth statement.",
            "When a gate is removed in VG0-VG3, the associated BPW/KBL operator remains and the effective gate is one; the ablation therefore compares gating alone. Global-gate histograms, spatial-gate mean/standard-deviation distributions, brightness/contrast relationships, and over-processing failures combine full-validation statistics with manual review.",
            f"The formal interpretation of performance and gate behavior remains at {pending('VGUP_INTERPRETATION')}.",
        ],
        "tables": [7, 8],
        "figures": [
            (5, "原图、BPW、全局门控、KBL、空间门与最终 VGUP 输出", "Input, BPW, global gating, KBL, spatial gate, and final VGUP output"),
        ],
    },
    {
        "number": "4.9",
        "zh": "跨数据集评价",
        "en": "Cross-Dataset Evaluation",
        "zh_paragraphs": [
            "S0 与 S1 在第二数据集上分别从官方 yolo11n.pt 独立训练，采用一致的输入、epoch、batch 和模型选择规则。主数据集训练权重到第二数据集、以及反向的零样本测试仅作为可选迁移评价，不能替代独立训练。",
            f"数据集差异、定量结果和场景案例待选定数据集后填入：{pending('CROSS_DATASET_INTERPRETATION')}。",
        ],
        "en_paragraphs": [
            "S0 and S1 are independently trained on the second dataset from official yolo11n.pt under the same input size, epoch budget, batch size, and selection rule. Optional zero-shot transfer from the primary dataset to the second dataset, and the reverse direction, cannot replace independent training.",
            f"Dataset differences, quantitative results, and scene-specific cases will be inserted after selection: {pending('CROSS_DATASET_INTERPRETATION')}.",
        ],
        "tables": [9],
        "figures": [],
    },
    {
        "number": "4.10",
        "zh": "跨模型泛化",
        "en": "Cross-Model Generalization",
        "zh_paragraphs": [
            "最终方法被适配到 YOLOv8n，以检验组合是否依赖 YOLO11n。适配保留 YOLOv8 的 C2f 结构和真实通道缩放，只将检测尺度移至 P2/P3/P4、使用 DySample，并在 Detect 前加入 CA-SCAM、输入端加入 VGUP。",
            f"M0/M1 与 YOLO11n 基线/最终结构的结果见表 10，正式解释占位为 {pending('CROSS_MODEL_INTERPRETATION')}。",
        ],
        "en_paragraphs": [
            "The final method is adapted to YOLOv8n to test whether the combination depends on YOLO11n. The adaptation retains YOLOv8-native C2f blocks and channel scaling, shifts detection to P2/P3/P4 with DySample, places CA-SCAM before Detect, and inserts VGUP at the input.",
            f"Table 10 compares M0/M1 with the YOLO11n baseline/final pair. Formal interpretation remains at {pending('CROSS_MODEL_INTERPRETATION')}.",
        ],
        "tables": [10],
        "figures": [],
    },
    {
        "number": "4.11",
        "zh": "模型复杂度与效率",
        "en": "Model Complexity and Efficiency",
        "zh_paragraphs": [
            "复杂度脚本统一记录输入、batch、warm-up、重复次数、FP32/FP16、是否包含预处理和 NMS、环境版本、延迟均值/标准差/P50/P95、FPS 与峰值显存。模型前向包含模型内部的 VGUP，但不默认包含图像解码、letterbox 和 NMS。",
            "精度-GFLOPs 与精度-Params 图使用所有具备必要字段的候选；精度-FPS 图按硬件、框架、精度和计时范围拆分，避免把不同口径画在同一速度坐标系。不会为了使最终模型处于有利区域而筛除模型。",
            f"正式权衡结论占位为 {pending('COMPLEXITY_INTERPRETATION')}。",
        ],
        "en_paragraphs": [
            "The complexity protocol records input size, batch size, warm-up and repeat counts, FP32/FP16 mode, preprocessing/NMS scope, environment versions, mean/std/P50/P95 latency, FPS, and peak memory. Model forward includes an in-model VGUP when present but excludes image decoding, letterbox preprocessing, and NMS by default.",
            "Accuracy-GFLOPs and accuracy-parameter plots include every candidate with the required fields. Accuracy-FPS plots are separated by hardware, framework, precision, and timing scope. Models are never removed merely because they are unfavorable to the proposed method.",
            f"The formal trade-off interpretation remains at {pending('COMPLEXITY_INTERPRETATION')}.",
        ],
        "tables": [11],
        "figures": [
            (6, "mAP50-95 与 GFLOPs/Params 的权衡", "mAP50-95 versus GFLOPs/parameters"),
            (7, "统一推理口径下 mAP50-95 与 FPS 的权衡", "mAP50-95 versus FPS under a unified inference protocol"),
        ],
    },
    {
        "number": "4.12",
        "zh": "定性与可解释性分析",
        "en": "Qualitative and Interpretability Analysis",
        "zh_paragraphs": [
            "代表性案例由 R00 与 R10 的逐图 TP/FP/FN 和匹配 IoU 差异自动生成候选清单，再由人工确认。候选覆盖基线漏检而最终模型检出、误检消除、定位改善、极小船、密集船、低对比度/低照度统计候选，以及最终模型失败案例。",
            "所有模型使用相同图像、置信度阈值和显示范围。检测框来自真实预测，不人工修改；特征图、CA-SCAM 校准图和 VGUP 门控图来自真实 forward hook。不会使用 AI 生成遥感图或伪造热力图。",
        ],
        "en_paragraphs": [
            "Representative candidates are ranked from per-image TP/FP/FN and matched-IoU differences between R00 and R10, followed by manual confirmation. The pool covers baseline misses recovered by the final model, removed false positives, localization improvements, tiny ships, dense ships, objective low-contrast/low-illumination candidates, and final-model failures.",
            "Every model uses the same image, confidence threshold, and display range. Boxes are real predictions and are not manually altered. Feature responses, CA-SCAM calibration maps, and VGUP gate maps are obtained from real forward hooks. AI-generated remote-sensing images and fabricated heatmaps are prohibited.",
        ],
        "tables": [],
        "figures": [
            (8, "统一阈值下的典型成功与失败案例", "Representative successes and failures under a common threshold"),
        ],
    },
    {
        "number": "4.13",
        "zh": "稳定性与失败案例分析",
        "en": "Stability and Failure-Case Analysis",
        "zh_paragraphs": [
            "R00、R02 与 R10 采用 seeds 0、1、2 报告 Recall、mAP50 和 mAP50-95 的 mean±std。其余消融先完成 seed 0，只有在结论敏感时再追加随机种子。",
            "失败分析预先覆盖极小目标、完全遮挡、厚云、岸线误检、海浪纹理、VGUP 过度处理、CA-SCAM 放大局部噪声和 DySample 训练波动。没有数据时只保留评价计划，不将可能性写成已观察结论。",
            f"稳定性与失败模式的正式结果占位为 {pending('STABILITY_FAILURE_INTERPRETATION')}。",
        ],
        "en_paragraphs": [
            "R00, R02, and R10 use seeds 0, 1, and 2 to report mean±std for Recall, mAP50, and mAP50-95. Other ablations initially use seed 0 and receive additional seeds only when the conclusion is sensitive.",
            "The preregistered failure analysis covers extremely small targets, full occlusion, thick clouds, shoreline false positives, wave texture, VGUP over-processing, CA-SCAM amplification of local noise, and DySample training variation. Until evidence exists, these remain evaluation categories rather than observed conclusions.",
            f"Formal stability and failure-mode results remain at {pending('STABILITY_FAILURE_INTERPRETATION')}.",
        ],
        "tables": [12],
        "figures": [
            (9, "多随机种子分布与失败案例分类", "Multi-seed distributions and categorized failure cases"),
        ],
    },
]


def _safe_output(path: Path, overwrite: bool) -> None:
    if path.exists() and not overwrite:
        raise FileExistsError(
            f"Refusing to overwrite existing manuscript file: {path}"
        )


def _md_table(table: dict[str, Any], language: str) -> str:
    headers = table[f"headers_{language}"]
    rows = table["rows"]
    lines = [
        "| " + " | ".join(headers) + " |",
        "|" + "|".join("---" for _ in headers) + "|",
    ]
    lines.extend("| " + " | ".join(map(str, row)) + " |" for row in rows)
    return "\n".join(lines)


def build_markdown(language: str) -> str:
    is_zh = language == "zh"
    title = "4. 实验与结果" if is_zh else "4. Experiments and Results"
    note = (
        "状态说明：本文件是正式实验章节初稿。所有 `{{PENDING_...}}` 均须由真实 run manifest、results.csv 或已验证统计填充；当前不预写方向性结论。"
        if is_zh
        else "Status note: this is a formal chapter draft. Every `{{PENDING_...}}` item must be filled from a real run manifest, results.csv, or validated statistic; no directional conclusion is pre-authored."
    )
    lines = [f"# {title}", "", note, ""]
    for section in SECTIONS:
        heading = section["zh"] if is_zh else section["en"]
        paragraphs = section["zh_paragraphs"] if is_zh else section["en_paragraphs"]
        lines += [f"## {section['number']} {heading}", ""]
        for paragraph in paragraphs:
            lines += [paragraph, ""]
        for formula in section.get("formulas", []):
            lines += [f"\\[{formula}\\]", ""]
        for table_number in section["tables"]:
            table = TABLES[table_number]
            caption = table["zh"] if is_zh else table["en"]
            prefix = "表" if is_zh else "Table"
            lines += [f"**{prefix} {table_number}. {caption}.**", "", _md_table(table, language), ""]
        for figure_number, zh_caption, en_caption in section["figures"]:
            caption = zh_caption if is_zh else en_caption
            prefix = "图" if is_zh else "Figure"
            lines += [
                f"> [{prefix} {figure_number} placeholder: {caption}; source must be a real dataset image, prediction, statistic, or forward hook.]",
                "",
                f"**{prefix} {figure_number}. {caption}.**",
                "",
            ]
    return "\n".join(lines)


def _latex_escape(value: Any) -> str:
    text = str(value)
    replacements = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "±": r"$\pm$",
    }
    return "".join(replacements.get(character, character) for character in text)


def build_latex() -> str:
    lines = [
        r"\documentclass[11pt]{article}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage{booktabs,longtable,array,amsmath,graphicx}",
        r"\begin{document}",
        r"\section{Experiments and Results}",
        r"\textbf{Draft status.} Every \texttt{\{\{PENDING\_...\}\}} item must be filled from real experiment artifacts; no directional conclusion is pre-authored.",
    ]
    for section in SECTIONS:
        lines.append(rf"\subsection{{{_latex_escape(section['en'])}}}")
        lines.extend(_latex_escape(item) for item in section["en_paragraphs"])
        for formula in section.get("formulas", []):
            lines.append(r"\[" + formula + r"\]")
        for table_number in section["tables"]:
            table = TABLES[table_number]
            columns = len(table["headers_en"])
            lines += [
                rf"\begin{{longtable}}{{{'l' * columns}}}",
                rf"\caption{{{_latex_escape(table['en'])}.}}\\",
                r"\toprule",
                " & ".join(_latex_escape(value) for value in table["headers_en"]) + r"\\",
                r"\midrule",
                r"\endfirsthead",
                r"\toprule",
                " & ".join(_latex_escape(value) for value in table["headers_en"]) + r"\\",
                r"\midrule",
                r"\endhead",
            ]
            lines.extend(
                " & ".join(_latex_escape(value) for value in row) + r"\\"
                for row in table["rows"]
            )
            lines += [r"\bottomrule", r"\end{longtable}"]
        for figure_number, _zh, en in section["figures"]:
            lines += [
                r"\begin{figure}[htbp]",
                r"\centering",
                rf"\fbox{{\parbox[c][1.2in][c]{{0.9\linewidth}}{{\centering Figure {figure_number} placeholder: {_latex_escape(en)}. Real evidence only.}}}}",
                rf"\caption{{{_latex_escape(en)}.}}",
                r"\end{figure}",
            ]
    lines += [r"\end{document}", ""]
    return "\n\n".join(lines)


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int], total: int = 9360) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if sum(widths) != total:
        raise ValueError((widths, sum(widths), total))
    table.autofit = False
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(total))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for value in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(value))
        grid.append(column)
    for row in table.rows:
        for cell, value in zip(row.cells, widths, strict=True):
            properties = cell._tc.get_or_add_tcPr()
            cell_width = properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                properties.append(cell_width)
            cell_width.set(qn("w:w"), str(value))
            cell_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _set_repeat_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _set_run_font(run, font: str, east_asia: str, size, **kwargs) -> None:
    from docx.oxml.ns import qn

    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = size
    for name, value in kwargs.items():
        if name == "color":
            run.font.color.rgb = value
        elif name == "bold":
            run.bold = value
        elif name == "italic":
            run.italic = value
        else:
            setattr(run.font, name, value)


def _add_page_field(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def _paragraph_border(paragraph, color="B7C4D2", fill=None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), "8")
        edge.set(qn("w:space"), "6")
        edge.set(qn("w:color"), color)
        borders.append(edge)
    properties.append(borders)
    if fill:
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), fill)
        properties.append(shading)


def _column_widths(headers: list[str]) -> list[int]:
    weights = []
    for header in headers:
        length = max(3, min(16, len(header)))
        weights.append(length)
    raw = [int(9360 * value / sum(weights)) for value in weights]
    raw[-1] += 9360 - sum(raw)
    return raw


def build_docx(language: str, output: Path) -> None:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor

    is_zh = language == "zh"
    body_font = "Calibri"
    east_asia = "宋体" if is_zh else "Arial"
    heading_east = "黑体" if is_zh else "Arial"
    blue = RGBColor(0x2E, 0x74, 0xB5)
    dark_blue = RGBColor(0x1F, 0x4D, 0x78)
    navy = RGBColor(0x20, 0x37, 0x48)
    muted = RGBColor(0x62, 0x6B, 0x73)
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = document.styles["Normal"]
    normal.font.name = body_font
    normal._element.rPr.rFonts.set(qn("w:ascii"), body_font)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), body_font)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    for style_name, size, color, before, after in (
        ("Heading 1", 16, blue, 18, 10),
        ("Heading 2", 13, blue, 12, 6),
        ("Heading 3", 12, dark_blue, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), heading_east)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    header_text = (
        "Ocean Engineering 论文初稿 | 实验与结果"
        if is_zh
        else "Ocean Engineering Manuscript Draft | Experiments and Results"
    )
    run = header.add_run(header_text)
    _set_run_font(run, "Arial", heading_east, Pt(8.5), color=muted)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    footer_run = footer.add_run("Page ")
    _set_run_font(footer_run, "Arial", heading_east, Pt(8.5), color=muted)
    _add_page_field(footer)

    # editorial_cover header pattern with a restrained academic override.
    for _ in range(3):
        spacer = document.add_paragraph()
        spacer.paragraph_format.space_after = Pt(18)
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker_run = kicker.add_run(
        "OCEAN ENGINEERING MANUSCRIPT DRAFT"
        if not is_zh
        else "OCEAN ENGINEERING 论文初稿"
    )
    _set_run_font(kicker_run, "Arial", heading_east, Pt(10), bold=True, color=blue)
    kicker.paragraph_format.space_after = Pt(14)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(
        "4. 实验与结果" if is_zh else "4. Experiments and Results"
    )
    _set_run_font(title_run, "Arial", heading_east, Pt(26), bold=True, color=navy)
    title.paragraph_format.space_after = Pt(8)
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "遥感小目标船舶检测正式实验章节骨架"
        if is_zh
        else "Formal Experiment Chapter Skeleton for Remote-Sensing Small-Ship Detection"
    )
    _set_run_font(subtitle_run, "Arial", heading_east, Pt(13), color=dark_blue)
    subtitle.paragraph_format.space_after = Pt(26)
    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata_run = metadata.add_run(
        "Generated from paper/formal-experiments | Values pending real runs"
    )
    _set_run_font(metadata_run, "Arial", heading_east, Pt(9.5), italic=True, color=muted)
    metadata.paragraph_format.space_after = Pt(80)
    policy = document.add_paragraph()
    policy.alignment = WD_ALIGN_PARAGRAPH.CENTER
    policy.paragraph_format.space_before = Pt(8)
    policy.paragraph_format.space_after = Pt(12)
    policy_run = policy.add_run(
        "所有 {{PENDING_...}} 仅可由真实实验文件填充；当前不包含推测性结果。"
        if is_zh
        else "Every {{PENDING_...}} item must come from a real experiment artifact; no result is inferred."
    )
    _set_run_font(policy_run, body_font, east_asia, Pt(10), bold=True, color=dark_blue)
    _paragraph_border(policy, fill="F4F6F9")
    document.add_page_break()

    for item in SECTIONS:
        heading = item["zh"] if is_zh else item["en"]
        paragraph = document.add_paragraph(
            f"{item['number']} {heading}",
            style="Heading 1",
        )
        paragraph.paragraph_format.keep_with_next = True
        paragraphs = item["zh_paragraphs"] if is_zh else item["en_paragraphs"]
        for text in paragraphs:
            body = document.add_paragraph()
            body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            body.add_run(text)
        for formula in item.get("formulas", []):
            equation = document.add_paragraph()
            equation.alignment = WD_ALIGN_PARAGRAPH.CENTER
            equation.paragraph_format.space_before = Pt(6)
            equation.paragraph_format.space_after = Pt(10)
            equation_run = equation.add_run(formula)
            _set_run_font(equation_run, "Cambria Math", "Cambria Math", Pt(11), italic=True)
        for table_number in item["tables"]:
            table_spec = TABLES[table_number]
            caption_text = table_spec["zh"] if is_zh else table_spec["en"]
            caption = document.add_paragraph()
            caption.paragraph_format.space_before = Pt(8)
            caption.paragraph_format.space_after = Pt(4)
            caption.paragraph_format.keep_with_next = True
            caption_run = caption.add_run(
                ("表" if is_zh else "Table")
                + f" {table_number}. {caption_text}."
            )
            _set_run_font(caption_run, body_font, east_asia, Pt(9.5), bold=True)
            headers = table_spec["headers_zh"] if is_zh else table_spec["headers_en"]
            table = document.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            widths = _column_widths(headers)
            _set_table_geometry(table, widths)
            _set_repeat_header(table.rows[0])
            font_size = Pt(7 if len(headers) >= 9 else 8.5)
            for index, value in enumerate(headers):
                cell = table.rows[0].cells[index]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                _set_cell_shading(cell, "F4F6F9")
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run(str(value))
                _set_run_font(run, "Arial", heading_east, font_size, bold=True, color=navy)
            for source_row in table_spec["rows"]:
                row = table.add_row()
                for index, value in enumerate(source_row):
                    cell = row.cells[index]
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.LEFT
                        if index in {0, 1} and len(headers) <= 7
                        else WD_ALIGN_PARAGRAPH.CENTER
                    )
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    run = paragraph.add_run(str(value))
                    _set_run_font(run, body_font, east_asia, font_size)
            after = document.add_paragraph()
            after.paragraph_format.space_after = Pt(2)
        for figure_number, zh_caption, en_caption in item["figures"]:
            caption_text = zh_caption if is_zh else en_caption
            placeholder = document.add_paragraph()
            placeholder.alignment = WD_ALIGN_PARAGRAPH.CENTER
            placeholder.paragraph_format.space_before = Pt(12)
            placeholder.paragraph_format.space_after = Pt(4)
            placeholder.paragraph_format.keep_with_next = True
            placeholder_run = placeholder.add_run(
                (
                    f"[图 {figure_number} 占位：{caption_text}；仅允许真实数据图、预测图、统计图或 forward hook。]"
                    if is_zh
                    else f"[Figure {figure_number} placeholder: {caption_text}; real data, predictions, statistics, or forward hooks only.]"
                )
            )
            _set_run_font(placeholder_run, body_font, east_asia, Pt(9.5), italic=True, color=muted)
            _paragraph_border(placeholder, fill="F8FAFC")
            caption = document.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.paragraph_format.space_after = Pt(8)
            caption_run = caption.add_run(
                ("图" if is_zh else "Figure")
                + f" {figure_number}. {caption_text}."
            )
            _set_run_font(caption_run, body_font, east_asia, Pt(9.5), bold=True)
    core = document.core_properties
    core.title = (
        "实验与结果章节骨架"
        if is_zh
        else "Experiments and Results Chapter Skeleton"
    )
    core.subject = "Ocean Engineering remote-sensing ship detection"
    core.author = "ship-yolo formal experiment framework"
    document.save(output)


def figure_table_inventory() -> str:
    lines = [
        "# 实验图表清单",
        "",
        "所有图表必须来自真实数据、正式 run、已验证统计或 forward hook；不得使用 AI 生成遥感图、伪造热力图或人工修改检测框。",
        "",
        "## 表格",
        "",
    ]
    for number, table in TABLES.items():
        lines.append(f"- 表 {number} / Table {number}: {table['zh']} / {table['en']}")
    lines += ["", "## 图", ""]
    figures = sorted(
        {
            number: (zh, en)
            for section in SECTIONS
            for number, zh, en in section["figures"]
        }.items()
    )
    for number, (zh, en) in figures:
        lines.append(f"- 图 {number} / Figure {number}: {zh} / {en}")
    lines += [
        "",
        "## 生成接口",
        "",
        "- DPLS: `evaluate_by_short_side.py`, `evaluate_by_area.py`, `analyze_detection_level_contribution.py`, `visualize_pyramid_features.py`, `compare_dpls_predictions.py`.",
        "- CA-SCAM: `visualize_ca_scam_forward.py`（真实 P2/P3/P4 前向 hook）.",
        "- VGUP: `visualize_vgup_forward.py`, `analyze_vgup_gates.py`.",
        "- 代表案例: `select_representative_cases.py`，自动结果只作为人工终选候选。",
        "- Complexity: `benchmark_model_complexity.py`, `benchmark_inference_latency.py`, `plot_complexity_tradeoff.py`.",
        "",
    ]
    return "\n".join(lines)


def placeholder_payload() -> dict[str, Any]:
    return {
        "schema_version": 1,
        "policy": {
            "replace_only_from_real_artifacts": True,
            "directional_conclusions_pre_authored": False,
            "keep_unresolved": True,
        },
        "format": "{{PENDING_{RUN_ID}_{FIELD}}}",
        "runs": RUNS,
        "run_fields": [
            "PRECISION",
            "RECALL",
            "MAP50",
            "MAP75",
            "MAP5095",
            "PARAMS",
            "GFLOPS",
            "LATENCY",
            "FPS",
        ],
        "additional_placeholders": sorted(
            {
                token
                for text in (
                    build_markdown("zh") + build_markdown("en")
                ).split()
                for token in [text.strip("`.,;:()[]")]
                if token.startswith("{{PENDING_") and token.endswith("}}")
            }
        ),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--output-dir",
        default=r"D:\遥感船舶检测论文\01_论文初稿",
    )
    parser.add_argument("--overwrite", action="store_true")
    args = parser.parse_args()
    output = Path(args.output_dir)
    output.mkdir(parents=True, exist_ok=True)
    files = {
        "zh_md": output / "04_实验与结果_中文.md",
        "en_md": output / "04_Experiments_and_Results_English.md",
        "tex": output / "04_Experiments_and_Results.tex",
        "zh_docx": output / "04_实验与结果_中文.docx",
        "en_docx": output / "04_Experiments_and_Results_English.docx",
        "inventory": output / "实验图表清单.md",
        "placeholders": output / "实验数据占位符.yaml",
    }
    for path in files.values():
        _safe_output(path, args.overwrite)
    files["zh_md"].write_text(build_markdown("zh"), encoding="utf-8")
    files["en_md"].write_text(build_markdown("en"), encoding="utf-8")
    files["tex"].write_text(build_latex(), encoding="utf-8")
    files["inventory"].write_text(figure_table_inventory(), encoding="utf-8")
    # JSON is a valid YAML 1.2 document and keeps this document builder free
    # from a PyYAML dependency in the bundled desktop document runtime.
    files["placeholders"].write_text(
        json.dumps(placeholder_payload(), ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    build_docx("zh", files["zh_docx"])
    build_docx("en", files["en_docx"])
    print(json.dumps({key: str(value) for key, value in files.items()}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
